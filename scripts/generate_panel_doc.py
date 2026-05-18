from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(31, 90, 69)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, center=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.1


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
    else:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(10.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        run.font.size = Pt(10.5)


def add_key_value_table(doc, rows, col_widths=(2.0, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    style_table(table)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(col_widths[0])
        cells[1].width = Inches(col_widths[1])
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
    return table


def add_matrix_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        if widths and i < len(widths):
            hdr[i].width = Inches(widths[i])
        set_cell_shading(hdr[i], "1F5A45")
        set_cell_text(hdr[i], head, bold=True, center=True, color=RGBColor(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if widths and i < len(widths):
                cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], val, center=(i > 0 and len(str(val)) <= 14))
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FurniComply Information Security Documentation")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Panel Format Draft")
    r.italic = True
    r.font.size = Pt(11)

    add_heading(doc, "Project Overview", 1)
    add_paragraph(doc, "This project is developed in partial fulfillment of the requirements for IT 16/L – Information Security 1.")
    add_paragraph(doc, "This documentation presents the design, implementation, and security considerations of the proposed system.")
    add_key_value_table(
        doc,
        [
            ("Prepared by", "[Your Name]"),
            ("Submitted to", "Cyril Loyd Tomas"),
            ("System Name", "FurniComply"),
            ("Course", "IT 16/L – Information Security 1"),
        ],
    )

    add_heading(doc, "System Description", 1)
    add_paragraph(
        doc,
        "FurniComply is a web-based compliance and procurement management system for furniture-industry operations. "
        "It is designed to manage and monitor regulatory compliance, supplier information, procurement records, "
        "corrective actions, and audit trails in one integrated platform."
    )
    add_paragraph(doc, "Core functionality includes:")
    add_bullets(
        doc,
        [
            "Policy governance for creating, approving, versioning, and archiving policies.",
            "Compliance checks for recording inspections, findings, and risk levels.",
            "CAPA management for assigning, tracking, and closing corrective and preventive actions.",
            "Supplier management for onboarding, document uploads, approval, and compliance scoring.",
            "Procurement workflow for purchase orders with approval and compliance gates.",
            "Regulatory reporting, dashboards, logs, user management, and database backup support.",
        ],
    )

    add_heading(doc, "Platform and Technologies Used", 1)
    add_matrix_table(
        doc,
        ["Category", "Technology"],
        [
            ["Programming Language", "C# (.NET 8)"],
            ["Framework / Environment", "ASP.NET Core MVC, Entity Framework Core"],
            ["Database", "SQLite for localhost development"],
            ["Platform", "Web Application"],
            ["Authentication", "ASP.NET Core Identity with PBKDF2 password hashing"],
            ["Supporting Libraries", "SecurityCodeScan, DotNetEnv, QuestPDF, ClosedXML, QRCoder"],
            ["Version Control", "Git / GitHub"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Security Policies", 1)

    add_heading(doc, "Password Policy", 2)
    add_matrix_table(
        doc,
        ["Rule", "Implementation"],
        [
            ["Minimum length", "12 characters"],
            ["Uppercase letters", "Required"],
            ["Lowercase letters", "Required"],
            ["Numbers", "Required"],
            ["Symbols", "Required"],
            ["Password storage", "Passwords are hashed with PBKDF2 and are not stored in plain text."],
            ["Password reset", "Token-based reset through Forgot Password."],
            ["Password reuse", "Recent password reuse is blocked during reset."],
            ["Periodic password updates", "Recommended as organizational policy; forced scheduled rotation is not enabled in the current build."],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Login Attempt Policy", 2)
    add_matrix_table(
        doc,
        ["Rule", "Implementation"],
        [
            ["Failed attempt limit", "5 failed attempts"],
            ["Temporary lockout", "5-minute lockout after the limit is exceeded"],
            ["Scope", "Applies to password and 2FA failures"],
            ["User feedback", "Generic invalid login message to reduce account enumeration"],
            ["Bot protection", "Google reCAPTCHA v2 on login and forgot-password"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Data Handling Policy", 2)
    add_matrix_table(
        doc,
        ["Data Type / Rule", "Implementation"],
        [
            ["Passwords", "Stored only as PBKDF2 hashes"],
            ["Sensitive fields", "Selected fields such as supplier contact data and some log IP values use AES-256-GCM field encryption"],
            ["Configuration secrets", ".env files and key folders are gitignored and kept outside version control"],
            ["Uploaded documents", "Document access is served through authorized controller actions only"],
            ["Session protection", "HTTP-only authentication cookie with optional idle auto-logout setting"],
            ["Access restriction", "Only authorized users can view or modify protected data"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Access Control Policy", 2)
    add_matrix_table(
        doc,
        ["Rule", "Implementation"],
        [
            ["System configuration", "Restricted to SuperAdmin and Admin roles"],
            ["General access model", "Role-based access control using authorization policies"],
            ["Anonymous access", "Limited to login, forgot/reset password, request access, privacy, and error pages"],
            ["Override actions", "Admin and SuperAdmin share administrative functions, but override actions are restricted to SuperAdmin only"],
            ["Access logging", "Login attempts and business actions are recorded in separate log areas"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Logging and Monitoring Policy", 2)
    add_matrix_table(
        doc,
        ["Policy Item", "Implementation"],
        [
            ["Security logs", "Records login success, failed login attempts, lockouts, and logout events"],
            ["Audit logs", "Records transaction and general business actions such as create, edit, approve, archive, and override"],
            ["Error handling", "Application errors are logged while users are shown safe generic messages"],
            ["Monitoring practice", "Administrators and compliance roles review logs for suspicious activity and audit evidence"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "Incident Response Plan", 1)
    add_matrix_table(
        doc,
        ["Phase", "Action"],
        [
            ["Detection", "Security incidents are identified through system logs, alerts, monitoring tools, and suspicious user activity."],
            ["Reporting", "Incidents are reported immediately to the system administrator or responsible authority."],
            ["Response", "Immediate containment actions include account lockout, password reset, session revocation, and isolation of affected components if needed."],
            ["Recovery", "System functionality is restored while validating data integrity and reviewing backup options."],
            ["Review", "Post-incident analysis is performed to improve future policies, controls, and code fixes."],
        ],
        widths=[1.6, 5.4],
    )

    add_heading(doc, "Code Auditing and Security Review", 1)
    add_heading(doc, "Tool Used", 2)
    add_matrix_table(
        doc,
        ["Tool", "Purpose"],
        [
            ["SecurityCodeScan.VS2019", "Roslyn-based security analyzer for insecure code patterns"],
            [".NET built-in analyzers", "Compiler and quality/security-related static analysis"],
            ["dotnet list package --vulnerable", "Dependency vulnerability checking"],
            ["GitHub Actions Security Checks", "Automated CI build and security scan workflow"],
            ["Run-SecurityAudit.ps1", "Local audit script that ties the checks together and generates latest-report.txt"],
        ],
        widths=[2.5, 4.5],
    )
    add_paragraph(doc, "One-sentence explanation: The system uses SecurityCodeScan, built-in .NET analyzers, dependency vulnerability scanning, and an automated GitHub Actions workflow to audit the codebase.")

    add_heading(doc, "Usage", 2)
    add_bullets(
        doc,
        [
            "The project is built in Release mode to surface analyzer findings during code review.",
            "The Run-SecurityAudit.ps1 script is executed from the repository root to collect tool inventory, dependency audit, and build findings.",
            "GitHub Actions repeats the security checks automatically on push and pull request.",
        ],
    )

    add_heading(doc, "Findings", 2)
    add_matrix_table(
        doc,
        ["ID", "Finding", "Status"],
        [
            ["F-01", "Anonymous password reset exposure risk", "Resolved by token-based reset flow"],
            ["F-02", "Risky debug or backup endpoints", "Restricted or removed"],
            ["F-03", "Unsafe public access to supplier document paths", "Blocked; authorized controller download only"],
            ["F-04", "Duplicate checks on encrypted supplier contact fields", "Resolved in application-layer validation"],
            ["F-05", "Transitive package advisories in optional dependency chain", "Documented and monitored"],
        ],
        widths=[0.8, 4.2, 2.0],
    )

    add_heading(doc, "Fixes", 2)
    add_bullets(
        doc,
        [
            "Removed or restricted risky endpoints and strengthened RBAC policies.",
            "Separated Security Logs from Audit Logs to align with the security documentation requirement.",
            "Improved password reset flow, reCAPTCHA integration, and password-history enforcement.",
            "Added supplier uniqueness checks that still work even when protected fields are encrypted.",
        ],
    )

    add_heading(doc, "Proof", 2)
    add_matrix_table(
        doc,
        ["Evidence", "Location / Note"],
        [
            ["Audit script", "scripts/Run-SecurityAudit.ps1"],
            ["CI workflow", ".github/workflows/security.yml"],
            ["Audit report", "security-audit/latest-report.txt"],
            ["Screenshot 1", "Insert screenshot of Visual Studio analyzer or Error List here"],
            ["Screenshot 2", "Insert screenshot of GitHub Actions Security Checks run here"],
            ["Screenshot 3", "Insert screenshot of dotnet package vulnerability output here"],
        ],
        widths=[1.8, 5.2],
    )

    add_heading(doc, "Access Control (RBAC / ACL)", 1)
    add_heading(doc, "Intended Users", 2)
    add_matrix_table(
        doc,
        ["Role", "Responsibility"],
        [
            ["SuperAdmin", "All administrative functions plus exclusive authority to perform override actions"],
            ["Admin", "Same administrative functions as SuperAdmin except override actions"],
            ["Compliance Manager", "Policies, compliance checks, CAPA, and reports"],
            ["Department Head", "Department-scoped checks, policies, and oversight"],
            ["Procurement", "Suppliers, orders, and procurement operations"],
            ["Auditor", "Read-only review of records and logs"],
            ["Guest", "Unauthenticated access to public pages only"],
        ],
        widths=[2.0, 5.0],
    )

    add_heading(doc, "Access Control Matrix", 2)
    add_paragraph(
        doc,
        "The system implements Role-Based Access Control (RBAC) with clear separation of duties. "
        "Each role is granted only the permissions necessary for its responsibilities, helping enforce "
        "security, accountability, and controlled access to system functions."
    )
    add_matrix_table(
        doc,
        ["System Module / Function", "SuperAdmin", "Admin", "Compliance Manager", "Department Head", "Procurement", "Auditor"],
        [
            ["User and Role Management", "√", "√", "x", "x", "x", "x"],
            ["System Configuration", "√", "√", "x", "x", "x", "x"],
            ["Database Backups", "√", "√", "x", "x", "x", "x"],
            ["Policy Create / Edit", "√", "√", "√", "√*", "x", "R"],
            ["Policy Approval", "√", "√", "x", "x", "x", "x"],
            ["Compliance Checks", "√", "√", "√", "√*", "x", "R"],
            ["CAPA Management", "√", "√", "√", "√*", "x", "R"],
            ["Supplier Management", "√", "√", "R", "x", "√", "x"],
            ["Supplier Approval", "√", "√", "x", "x", "x", "x"],
            ["Supplier Override", "√", "x", "x", "x", "x", "x"],
            ["Procurement Orders", "√", "√", "R", "x", "√", "x"],
            ["Regulatory Reports", "√", "√", "√", "R", "x", "R"],
            ["Report Approval", "√", "√", "√", "x", "x", "x"],
            ["Report Override", "√", "x", "x", "x", "x", "x"],
            ["Audit Logs", "√", "√", "R", "O", "O", "R"],
            ["Security Logs", "√", "√", "R", "O", "O", "R"],
        ],
        widths=[2.4, 0.85, 0.85, 1.25, 1.2, 1.0, 0.85],
    )

    add_paragraph(doc, "Legend: √ = allowed, x = denied, R = read-only, O = own records / scoped view only, √* = allowed only within assigned department scope.", italic=True)

    doc.add_page_break()
    add_heading(doc, "Final Submission Notes", 1)
    add_bullets(
        doc,
        [
            "Replace [Your Name] before submission.",
            "Attach the required screenshots in the Proof section.",
            "If the panel provides a cover-page format, this content can be transferred into that template without changing the tables.",
        ],
    )

    return doc


if __name__ == "__main__":
    output_path = r"D:\Furniture\docs\Information-Security-Documentation-Panel-Format.docx"
    build_doc().save(output_path)
    print(output_path)
